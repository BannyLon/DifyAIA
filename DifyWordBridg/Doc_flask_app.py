from flask import Flask, request, jsonify
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import logging
import subprocess

app = Flask(__name__)

# 配置日志
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# 配置保存文档的目录
SAVE_DIR = "./documents/"
if not os.path.exists(SAVE_DIR):
    os.makedirs(SAVE_DIR)

@app.route('/generate_doc', methods=['POST'])
def generate_doc():
    try:
        # 获取请求中的JSON数据
        data = request.json
        title = data.get('title')
        content = data.get('content')

        if not title and not content:
            logger.error("Title or content is required")
            return jsonify({"error": "Title or content is required"}), 400

        # 生成文档
        file_name = f"phl_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(SAVE_DIR, file_name)
        logger.debug(f"File path: {file_path}")

        doc = Document()

        if title:
            # 添加大标题
            paragraph = doc.add_heading(title, level=1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
            paragraph.style.font.name = 'FangSong'  # 直接设置整个段落的字体
            paragraph.style.font.size = Pt(22)  # 二号字体

        if content:
            # 添加正文
            paragraph = doc.add_paragraph(content)
            paragraph.style.font.name = 'FangSong'  # 直接设置整个段落的字体
            paragraph.style.font.size = Pt(10.5)  # 五号字体

        doc.save(file_path)
        logger.info(f"Document generated successfully at {file_path}")

        # 在Mac上打开文件
        subprocess.call(['open', file_path])

        return jsonify({"message": "Document generated successfully", "file_path": file_path}), 200

    except Exception as e:
        logger.error(f"Error generating document: {e}")
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5001)